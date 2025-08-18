import re
import time
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


class AIProcessor:
    """AI处理模块 - 用于提炼和总结概念内容（本地模拟版本）"""
    
    def __init__(self):
        """
        初始化本地AI处理器
        """
        pass
        
    def prepare_content_for_ai(self, knowledge_tree):
        """
        准备发送给AI的内容
        :param knowledge_tree: 知识树结构
        :return: 格式化的内容字符串
        """
        content_parts = []
        for node in knowledge_tree:
            section_text = f"## {node['title']}\n"
            if isinstance(node['content'], list):
                section_text += "\n".join(node['content'])
            else:
                section_text += str(node['content'])
            content_parts.append(section_text)
        
        return "\n\n".join(content_parts)
    
    def send_to_ai(self, content, model=None):
        """
        使用本地模拟模型处理内容
        :param content: 要处理的内容
        :param model: 模型参数（本地处理不使用）
        :return: 处理结果
        """
        # 模拟处理时间
        time.sleep(1)
        
        # 调用本地处理函数
        return self._local_process(content)
    
    def _local_process(self, content):
        """
        本地内容处理函数
        :param content: 输入内容
        :return: 处理后的内容
        """
        # 按章节分割内容
        sections = content.split("##")
        processed_sections = []
        
        for section in sections:
            if section.strip():
                lines = section.strip().split("\n")
                title = lines[0].strip() if lines else "未命名章节"
                content_lines = lines[1:] if len(lines) > 1 else []
                
                # 提取关键要点
                key_points = self._extract_key_points(content_lines)
                
                # 格式化输出
                processed_content = "\n".join([f"- {point}" for point in key_points]) if key_points else "无具体内容"
                processed_sections.append(f"## {title}\n{processed_content}")
        
        return "\n\n".join(processed_sections)
    
    def _extract_key_points(self, content_lines):
        """
        从内容行中提取关键要点
        :param content_lines: 内容行列表
        :return: 关键要点列表
        """
        if not content_lines:
            return []
        
        # 合并所有内容行
        full_content = " ".join(content_lines)
        
        # 简单的关键词提取逻辑
        # 这里可以扩展为更复杂的自然语言处理逻辑
        sentences = re.split(r'[.。!?！？;；]', full_content)
        key_points = []
        
        # 提取包含重要词汇的句子作为要点
        important_words = ['重要', '关键', '核心', '主要', '首先', '其次', '最后', '总之', '因此', '所以', '特点', '优势', '定义', '概念']
        
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and any(word in sentence for word in important_words):
                # 简化句子，只保留主要信息
                simplified = self._simplify_sentence(sentence)
                if simplified and len(simplified) > 10:  # 过滤太短的句子
                    key_points.append(simplified)
        
        # 如果没有找到包含重要词汇的句子，则选择前几句话作为要点
        if not key_points:
            for i, sentence in enumerate(sentences[:3]):  # 取前3句
                sentence = sentence.strip()
                if sentence and len(sentence) > 10:
                    simplified = self._simplify_sentence(sentence)
                    if simplified:
                        key_points.append(simplified)
        
        return key_points[:5]  # 最多返回5个要点
    
    def _simplify_sentence(self, sentence):
        """
        简化句子，去除冗余信息
        :param sentence: 原始句子
        :return: 简化后的句子
        """
        # 去除多余空格
        sentence = re.sub(r'\s+', ' ', sentence).strip()
        
        # 去除一些常见的冗余表述
        redundant_patterns = [
            r'在这个.*?中',
            r'通过.*?可以发现',
            r'从.*?可以看出',
            r'根据.*?得知',
            r'由于.*?所以',
        ]
        
        for pattern in redundant_patterns:
            sentence = re.sub(pattern, '', sentence)
        
        return sentence.strip()
    
    def parse_ai_response(self, ai_response):
        """
        解析AI响应为结构化数据
        :param ai_response: AI响应文本
        :return: 结构化的知识树
        """
        # 按章节分割
        sections = ai_response.split("##")
        knowledge_tree = []
        
        for section in sections:
            if section.strip():
                lines = section.strip().split("\n")
                title = lines[0].strip() if lines else "未命名章节"
                content = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
                
                knowledge_tree.append({
                    'title': title,
                    'content': content,
                    'children': []
                })
        
        # 如果解析失败，返回原始响应
        if not knowledge_tree:
            knowledge_tree.append({
                'title': 'AI总结结果',
                'content': ai_response,
                'children': []
            })
            
        return knowledge_tree
    
    def export_to_word(self, knowledge_tree, output_path):
        """
        将知识树导出为Word文档
        :param knowledge_tree: 知识树结构
        :param output_path: 输出路径
        :return: 输出文件路径
        """
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 确保中文字体正确应用
        style.element.xpath('.//w:rPr')[0].insert(0, OxmlElement('w:rFonts'))
        style.element.xpath('.//w:rFonts')[0].set(qn('w:eastAsia'), '宋体')
        
        # 设置标题样式
        heading_style = doc.styles['Heading 1']
        heading_font = heading_style.font
        heading_font.name = '宋体'
        heading_font.size = Pt(16)
        
        heading_style.element.xpath('.//w:rPr')[0].insert(0, OxmlElement('w:rFonts'))
        heading_style.element.xpath('.//w:rFonts')[0].set(qn('w:eastAsia'), '宋体')
        
        for node in knowledge_tree:
            # 添加标题
            heading = doc.add_heading(node['title'], level=1)
            heading.style = 'Heading 1'
            for run in heading.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加内容
            content = node['content'] if isinstance(node['content'], str) else "\n".join(node['content'])
            if content.strip():
                paragraph = doc.add_paragraph(content)
                paragraph.style = 'Normal'
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.save(output_path)
        return output_path
    
    def export_to_pdf(self, word_path, pdf_path):
        """
        将Word文档转换为PDF
        :param word_path: Word文件路径
        :param pdf_path: PDF输出路径
        :return: PDF文件路径
        """
        try:
            import comtypes.client
            import os
            
            # 确保路径为绝对路径
            word_path = os.path.abspath(word_path)
            pdf_path = os.path.abspath(pdf_path)
            
            # 创建Word应用实例
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            
            # 打开Word文档
            doc = word.Documents.Open(word_path)
            
            # 保存为PDF
            doc.SaveAs2(pdf_path, FileFormat=17)  # 17是PDF格式代码
            
            # 关闭文档和Word进程
            doc.Close()
            word.Quit()
            
            return pdf_path if os.path.exists(pdf_path) else None
        except Exception as e:
            print(f"PDF导出失败: {str(e)}")
            return None